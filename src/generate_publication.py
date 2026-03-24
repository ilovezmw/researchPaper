"""
Build a research note .docx from structured YAML (matches the two-column + AI disclosure layout).

Usage:
  python src/generate_publication.py --input "In Process/my-topic"
  python src/generate_publication.py --input "In Process/my-topic" --pdf
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

# Allow running as script from repo root
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from src.docx_layout import (  # noqa: E402
    apply_document_defaults,
    configure_page,
    polish_research_table,
    set_section_columns,
    set_table_rows_cant_split,
    style_abstract_body,
    style_abstract_label,
    style_body,
    style_heading,
    style_meta_line,
    style_reference,
    style_section_divider_heading,
    style_table_caption,
    style_table_note,
    style_title,
    style_title_subline,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# Long single <w:t> runs can misbehave in some Word builds; chunk when adding body text.
_RUN_TEXT_CHUNK = 240


def resolve_published_dir(root: Path) -> Path:
    """Prefer Published/ if it already exists (matches common Windows folder naming)."""
    caps = root / "Published"
    low = root / "published"
    if caps.is_dir():
        return caps
    if low.is_dir():
        return low
    caps.mkdir(parents=True, exist_ok=True)
    return caps


def resolve_topic_output_dir(published_root: Path, topic_dir: Path) -> Path:
    """
    Mirror the topic folder name under Published/.
    Example:
      In Process/My Topic -> Published/My Topic/
    """
    out = published_root / topic_dir.name
    out.mkdir(parents=True, exist_ok=True)
    return out


def _docx_plain_text(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    parts: list[str] = []
    for t in root.iter(f"{{{_W_NS}}}t"):
        if t.text:
            parts.append(t.text)
        if t.tail:
            parts.append(t.tail)
    return "".join(parts)


def verify_docx_against_yaml(docx_path: Path, data: dict) -> list[str]:
    """Post-save checks so obviously truncated exports are visible in the terminal."""
    issues: list[str] = []
    blob = _docx_plain_text(docx_path)
    yaml_chars = sum(len(c) for c in _yaml_prose_paragraphs(data.get("abstract") or ""))
    para_n = 0
    for sec in data.get("sections") or []:
        for p in sec.get("paragraphs") or []:
            yaml_chars += sum(len(b) for b in _yaml_prose_paragraphs(str(p)))
            para_n += 1
    for sec in data.get("sections") or []:
        h = sec.get("heading")
        if h and str(h) not in blob:
            issues.append(f'missing section heading in DOCX: "{h}"')
    if yaml_chars >= 2500 and len(blob) < int(yaml_chars * 0.35):
        issues.append(
            f"DOCX text looks too short ({len(blob)} chars vs ~{yaml_chars} in YAML prose); "
            "regenerate after saving manuscript.yaml or close the .docx in Word if it locked an old export."
        )
    return issues


def _add_runs_chunked(paragraph, text: str) -> None:
    t = str(text)
    if not t:
        return
    for i in range(0, len(t), _RUN_TEXT_CHUNK):
        paragraph.add_run(t[i : i + _RUN_TEXT_CHUNK])


def _yaml_prose_paragraphs(text: str) -> list[str]:
    """
    YAML | block scalars keep newline characters; editors often break lines mid-sentence.
    - Split on blank lines for real paragraph breaks.
    - Within each chunk, join wrapped lines with a single space (fixes 'liquidity' + 'cycles' glue).
    - Collapse accidental multiple spaces.
    """
    raw = (text or "").strip()
    if not raw:
        return []
    out: list[str] = []
    for chunk in re.split(r"\n\s*\n", raw):
        lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
        if not lines:
            continue
        merged = " ".join(lines)
        merged = re.sub(r" +", " ", merged).strip()
        out.append(merged)
    return out


DEFAULT_AI_DISCLOSURE = (
    "Portions of this research note were developed with the assistance of AI-based analytical and editorial tools to "
    "support research synthesis, language refinement, and structural organization. All analytical frameworks, "
    "interpretations, and investment perspectives presented in this report are independently developed and remain "
    "the sole responsibility of the author."
)


def _slug(s: str) -> str:
    s = re.sub(r"[^\w\s-]", "", s, flags=re.UNICODE)
    s = re.sub(r"[-\s]+", "_", s.strip()).strip("_")
    return s or "research_note"


def _add_paragraph(document: Document, text: str, style_fn) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    style_fn(p)


def _add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.add_run(text)
    style_heading(p)


def _append_continuous_section(document: Document, num_cols: int) -> None:
    document.add_section(WD_SECTION.CONTINUOUS)
    s = document.sections[-1]
    configure_page(s)
    set_section_columns(s, num_cols)


def _apply_table_grid_fonts(table, font_pt: int = 8) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_pt)


def _add_data_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    col_width_in: float | None,
    *,
    font_pt: int = 9,
) -> None:
    ncols = len(headers)
    table = document.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    set_table_rows_cant_split(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.cell(ri + 1, ci).text = str(val)
    _apply_table_grid_fonts(table, font_pt=font_pt)
    if col_width_in is not None and ncols > 0:
        w = Inches(col_width_in)
        for row in table.rows:
            for cell in row.cells:
                cell.width = w
    polish_research_table(table)


def _column_break_paragraph(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    r.add_break(WD_BREAK.COLUMN)


def _emit_table_blocks(
    document: Document,
    tables: list[dict],
    *,
    default_layout: str,
) -> None:
    """
    full_page: insert a single-column section so the table is not split across columns (table uses full text width).
    narrow: keep inside the current two-column section; use column break + cantSplit to reduce splits.
    """
    i = 0
    while i < len(tables):
        tbl = tables[i]
        layout = (tbl.get("layout") or default_layout or "full_page").strip().lower()
        if layout == "narrow":
            if tbl.get("column_break_before", True):
                _column_break_paragraph(document)
            cap = tbl.get("caption")
            if cap:
                p = document.add_paragraph()
                p.add_run(str(cap))
                style_table_caption(p)
                p.paragraph_format.keep_with_next = True
            headers = tbl.get("headers") or []
            rows = tbl.get("rows") or []
            _add_data_table(
                document,
                headers,
                rows,
                tbl.get("col_width_in"),
                font_pt=9,
            )
            note = tbl.get("note")
            if note:
                p = document.add_paragraph()
                _add_runs_chunked(p, str(note))
                style_table_note(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i += 1
            continue

        # Batch consecutive full-page tables into one single-column section
        group: list[dict] = [tbl]
        j = i + 1
        while j < len(tables):
            nxt = tables[j]
            nxt_layout = (nxt.get("layout") or default_layout or "full_page").strip().lower()
            if nxt_layout == "narrow":
                break
            group.append(nxt)
            j += 1

        _append_continuous_section(document, 1)
        for t in group:
            cap = t.get("caption")
            if cap:
                p = document.add_paragraph()
                p.add_run(str(cap))
                style_table_caption(p)
                p.paragraph_format.keep_with_next = True
            headers = t.get("headers") or []
            rows = t.get("rows") or []
            _add_data_table(
                document,
                headers,
                rows,
                t.get("col_width_in"),
                font_pt=9,
            )
            note = t.get("note")
            if note:
                p = document.add_paragraph()
                _add_runs_chunked(p, str(note))
                style_table_note(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _append_continuous_section(document, 2)
        i = j


def build_document(data: dict) -> Document:
    document = Document()
    apply_document_defaults(document)

    # --- Section 0: title block + abstract (single column) ---
    s0 = document.sections[0]
    configure_page(s0)
    set_section_columns(s0, 1)

    title = data.get("title", "Untitled")
    title_lines = title if isinstance(title, list) else [title]
    for idx, line in enumerate(title_lines):
        p = document.add_paragraph()
        p.add_run(str(line))
        if idx == 0:
            style_title(p)
        else:
            style_title_subline(p)

    meta = [
        ("author", False),
        ("affiliation", False),
        ("location", False),
        ("date", False),
    ]
    for key, bold in meta:
        val = data.get(key)
        if not val:
            continue
        p = document.add_paragraph()
        p.add_run(str(val))
        style_meta_line(p, bold=bold)

    abstract_chunks = _yaml_prose_paragraphs(data.get("abstract") or "")
    if abstract_chunks:
        p = document.add_paragraph()
        p.add_run("Abstract")
        style_abstract_label(p)
        for block in abstract_chunks:
            p2 = document.add_paragraph()
            _add_runs_chunked(p2, block)
            style_abstract_body(p2)
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Body: two columns; tables default to a single-column "island" to avoid splitting across columns ---
    _append_continuous_section(document, 2)
    default_table_layout = (data.get("default_table_layout") or "full_page").strip().lower()

    for sec in data.get("sections") or []:
        heading = sec.get("heading")
        if heading:
            _add_heading(document, str(heading))
        for para in sec.get("paragraphs") or []:
            for block in _yaml_prose_paragraphs(str(para)):
                p = document.add_paragraph()
                _add_runs_chunked(p, block)
                style_body(p)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        tables = sec.get("tables") or []
        if tables:
            _emit_table_blocks(document, tables, default_layout=default_table_layout)

    # --- Final section: AI disclosure + references (single column) ---
    document.add_section(WD_SECTION.CONTINUOUS)
    s_ai = document.sections[-1]
    configure_page(s_ai)
    set_section_columns(s_ai, 1)

    p = document.add_paragraph()
    p.add_run("AI Assistance Disclosure")
    style_section_divider_heading(p)

    disclosure_src = (data.get("ai_disclosure") or DEFAULT_AI_DISCLOSURE).strip()
    disc_paras = _yaml_prose_paragraphs(disclosure_src)
    for db in disc_paras:
        p2 = document.add_paragraph()
        _add_runs_chunked(p2, db)
        style_body(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    refs = data.get("references") or []
    if refs:
        pr = document.add_paragraph()
        pr.add_run("References")
        style_section_divider_heading(pr)
        for ref in refs:
            rp = document.add_paragraph()
            _add_runs_chunked(rp, str(ref))
            style_reference(rp)

    return document


def load_yaml(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    return yaml.safe_load(text) or {}


def export_pdf_word_win32(docx_path: Path, pdf_path: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as e:
        raise RuntimeError("PDF export on Windows requires pywin32. Install: pip install pywin32") from e

    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
    finally:
        pythoncom.CoUninitialize()


def mark_done(topic_dir: Path) -> None:
    marker = topic_dir / "Done.txt"
    marker.write_text(
        "Marked complete by generate_publication.py — you can remove this file if you prefer another convention.\n",
        encoding="utf-8",
    )


def generate_one_topic(
    *,
    root: Path,
    topic_dir: Path,
    output_stem_override: str | None,
    export_pdf: bool,
    mark_done_file: bool,
) -> int:
    manuscript = topic_dir / "manuscript.yaml"
    if not manuscript.is_file():
        print(f"Missing manuscript.yaml in {topic_dir}", file=sys.stderr)
        return 2

    data = load_yaml(manuscript)
    doc = build_document(data)
    published_root = resolve_published_dir(root)
    topic_output_dir = resolve_topic_output_dir(published_root, topic_dir)

    if output_stem_override:
        base = _slug(output_stem_override)
    else:
        yaml_stem = data.get("output_stem")
        if yaml_stem:
            base = _slug(str(yaml_stem))
        else:
            title_slug = _slug(str(data.get("title") if not isinstance(data.get("title"), list) else data["title"][0]))
            base = f"{title_slug}_{data.get('date', '').replace(' ', '_')}"
    docx_name = f"{base}.docx"
    docx_path = (topic_output_dir / docx_name).resolve()

    doc.save(str(docx_path))
    print(f"Project root:  {root.resolve()}")
    print(f"Manuscript:      {manuscript.resolve()}")
    print(f"Wrote DOCX:      {docx_path}")
    for msg in verify_docx_against_yaml(docx_path, data):
        print(f"WARNING: {msg}", file=sys.stderr)

    if export_pdf:
        pdf_path = (topic_output_dir / f"{base}.pdf").resolve()
        export_pdf_word_win32(docx_path, pdf_path)
        print(f"Wrote PDF:       {pdf_path}")

    if mark_done_file:
        mark_done(topic_dir)
        print(f"Marked done: {topic_dir / 'Done.txt'}")

    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate research note DOCX (and optional PDF) from manuscript.yaml")
    parser.add_argument("--input", "-i", required=True, help="Path to topic folder under In Process")
    parser.add_argument(
        "--output-stem",
        "-n",
        default=None,
        metavar="NAME",
        help="Output file base name (no extension). Use when several folders share the same YAML title, or to avoid overwrites.",
    )
    parser.add_argument("--pdf", action="store_true", help="Also export PDF via Microsoft Word (Windows)")
    parser.add_argument("--no-done", action="store_true", help="Do not write Done.txt in the topic folder")
    parser.add_argument(
        "--batch",
        action="store_true",
        help="Treat --input as a parent folder and process each immediate subfolder that contains manuscript.yaml",
    )
    args = parser.parse_args()

    root = Path(__file__).resolve().parent.parent
    topic_dir = (root / args.input).resolve() if not Path(args.input).is_absolute() else Path(args.input).resolve()
    if not topic_dir.is_dir():
        print(f"Input folder not found: {topic_dir}", file=sys.stderr)
        return 2

    if args.batch:
        topic_dirs = sorted(
            [
                p
                for p in topic_dir.iterdir()
                if p.is_dir() and (p / "manuscript.yaml").is_file()
            ],
            key=lambda p: p.name.lower(),
        )
        if not topic_dirs:
            print(f"No topic subfolders with manuscript.yaml under {topic_dir}", file=sys.stderr)
            return 2
        for td in topic_dirs:
            code = generate_one_topic(
                root=root,
                topic_dir=td,
                output_stem_override=args.output_stem,
                export_pdf=args.pdf,
                mark_done_file=not args.no_done,
            )
            if code != 0:
                return code
        return 0

    return generate_one_topic(
        root=root,
        topic_dir=topic_dir,
        output_stem_override=args.output_stem,
        export_pdf=args.pdf,
        mark_done_file=not args.no_done,
    )


if __name__ == "__main__":
    raise SystemExit(main())
