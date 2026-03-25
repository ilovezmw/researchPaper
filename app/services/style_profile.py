"""
从参考 DOCX 提取字体、字号、行距、页边距、分栏等「样式配置」。
若文件不存在或解析失败，回退到学术默认（Times New Roman 12pt、双栏正文等）。
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.section import Section


@dataclass
class StyleProfile:
    """可序列化的默认排版参数，供 formatter 统一应用。"""

    body_font_name: str = "Times New Roman"
    body_font_size_pt: float = 12.0
    heading_font_name: str = "Times New Roman"
    title_size_pt: float = 18.0
    h1_size_pt: float = 14.0
    h2_size_pt: float = 13.0
    line_spacing: float = 1.15
    # 页边距（英寸），常见双栏/课程论文约 1in
    margin_top_in: float = 1.0
    margin_bottom_in: float = 1.0
    margin_left_in: float = 1.0
    margin_right_in: float = 1.0
    # 正文栏数（≥2 为双栏/多栏）；栏间距为 Word twips（1/20 pt），常见约 708（~0.5"）
    body_columns: int = 2
    column_spacing_twips: int | None = 708
    # True：标题/作者区单栏，正文用连续分节符后改为 body_columns（常见 IEEE/会议模板）
    title_section_single_column: bool = True


def _parse_cols_from_sect_pr(sect_pr) -> tuple[int, int | None]:
    """从节的 sectPr 读取 w:cols：栏数与栏间距（twips）。"""
    cols_el = sect_pr.find(qn("w:cols"))
    if cols_el is None:
        return (1, None)
    num = cols_el.get(qn("w:num"))
    if not num:
        return (1, None)
    try:
        n = int(num)
    except ValueError:
        return (1, None)
    sp = cols_el.get(qn("w:space"))
    sp_i: int | None = int(sp) if sp else None
    return (n, sp_i)


def set_section_column_count(section: Section, num: int, space_twips: int | None = None) -> None:
    """
    设置节的栏数（OOXML w:cols）。num=1 为单栏；num>=2 为分栏。
    python-docx 无高层 API，直接写 sectPr。
    """
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:cols"))
    if old is not None:
        sect_pr.remove(old)

    cols_el = OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(max(1, num)))
    if num >= 2 and space_twips is not None and space_twips > 0:
        cols_el.set(qn("w:space"), str(int(space_twips)))
    sect_pr.append(cols_el)


def _first_body_font(doc: Document) -> tuple[str | None, float | None]:
    """取正文第一个非空段落的字体与字号作为参考。"""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        for run in p.runs:
            if run.text and run.text.strip():
                name = run.font.name
                size = run.font.size
                pt = size.pt if size is not None else None
                return (name, pt)
        return (None, None)
    return (None, None)


def load_profile_from_reference(reference_path: Path) -> StyleProfile:
    """
    从参考 DOCX 加载样式；文件不存在时返回默认 profile。
    说明：python-docx 对「样式表」支持有限，这里以首段 run 为主信号。
    """
    profile = StyleProfile()
    if not reference_path.is_file():
        return profile
    try:
        doc = Document(str(reference_path))
    except Exception:
        return profile

    name, pt = _first_body_font(doc)
    if name:
        profile.body_font_name = name
        profile.heading_font_name = name
    if pt is not None and 8 <= pt <= 18:
        profile.body_font_size_pt = float(pt)

    # 尝试从第一节取页边距
    try:
        sec = doc.sections[0]
        profile.margin_top_in = float(sec.top_margin / Inches(1))
        profile.margin_bottom_in = float(sec.bottom_margin / Inches(1))
        profile.margin_left_in = float(sec.left_margin / Inches(1))
        profile.margin_right_in = float(sec.right_margin / Inches(1))
    except Exception:
        pass

    # 首段行距（若可解析）
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        fmt = p.paragraph_format
        if fmt.line_spacing_rule == WD_LINE_SPACING.MULTIPLE and fmt.line_spacing:
            profile.line_spacing = float(fmt.line_spacing)
        elif fmt.line_spacing and fmt.line_spacing > 0:
            # 可能是绝对行距，保守保持默认
            pass
        break

    # 从各节 sectPr 推断分栏：取最大栏数与对应间距；多节且首节单栏、后续多栏则与「标题单栏+正文双栏」一致
    col_specs: list[tuple[int, int | None]] = []
    for sec in doc.sections:
        col_specs.append(_parse_cols_from_sect_pr(sec._sectPr))
    max_cols = max((c[0] for c in col_specs), default=1)
    if max_cols >= 2:
        profile.body_columns = max_cols
        for n, sp in col_specs:
            if n >= 2 and sp is not None:
                profile.column_spacing_twips = sp
                break
    if len(col_specs) >= 2 and col_specs[0][0] == 1 and any(c[0] >= 2 for c in col_specs[1:]):
        profile.title_section_single_column = True
    elif col_specs and col_specs[0][0] >= 2 and len(col_specs) == 1:
        # 整篇一节且已为双栏：不强制再拆标题节，与参考一致
        profile.title_section_single_column = False

    return profile


def apply_section_margins(doc: Document, profile: StyleProfile) -> None:
    """将边距应用到文档所有节。"""
    for sec in doc.sections:
        sec.top_margin = Inches(profile.margin_top_in)
        sec.bottom_margin = Inches(profile.margin_bottom_in)
        sec.left_margin = Inches(profile.margin_left_in)
        sec.right_margin = Inches(profile.margin_right_in)


def make_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
