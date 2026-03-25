"""
将任意「论文体」DOCX 按 StyleProfile 与章节启发式重排为统一学术版式。
限制：复杂域代码、脚注、尾注、图片仅尽力保留；表格通过 OOXML 深拷贝保留结构。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

from app.config import AUTHOR_LINE_1, AUTHOR_LINE_2, AUTHOR_LINE_3
from app.services.docx_parser import append_table_copy, blocks_to_plain_lines, iter_body_blocks, paragraph_text
from app.services.style_profile import (
    StyleProfile,
    apply_section_margins,
    make_run_font,
    set_section_column_count,
)

logger = logging.getLogger(__name__)

# 独立标题行匹配：整行仅为章节名（可带编号），用于切换章节而非当作正文
_HEADING_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("abstract", re.compile(r"^\s*\d*\.?\s*abstract\s*:?\s*$", re.I)),
    (
        "introduction",
        re.compile(
            r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)introduction\s*:?\s*$",
            re.I,
        ),
    ),
    (
        "related_work",
        re.compile(
            r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)(related\s+work|literature\s+review|background)\s*:?\s*$",
            re.I,
        ),
    ),
    (
        "method",
        re.compile(
            r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)(methodology|method|architecture|proposed\s+system|approach|design)\s*:?\s*$",
            re.I,
        ),
    ),
    (
        "experiments",
        re.compile(
            r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)(experiment|experiments|evaluation|results|implementation)\s*:?\s*$",
            re.I,
        ),
    ),
    (
        "discussion",
        re.compile(r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)discussion\s*:?\s*$", re.I),
    ),
    (
        "conclusion",
        re.compile(r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)conclusions?\s*:?\s*$", re.I),
    ),
    (
        "references",
        re.compile(
            r"^\s*(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d*\.?\s*)(references?|bibliography|works\s+cited)\s*:?\s*$",
            re.I,
        ),
    ),
]

_SECTION_DISPLAY = {
    "abstract": "Abstract",
    "introduction": "Introduction",
    "related_work": "Related Work",
    "method": "Method",
    "experiments": "Experiments",
    "discussion": "Discussion",
    "conclusion": "Conclusion",
    "references": "References",
}


def _normalize_heading_line(s: str) -> str:
    """去掉零宽字符、合并空白，避免 Word 里看不见字符导致标题匹配失败。"""
    for ch in ("\u200b", "\u200c", "\u200d", "\ufeff", "\xa0"):
        s = s.replace(ch, " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _match_executive_summary_heading(line: str) -> bool:
    s = _normalize_heading_line(line)
    return bool(re.match(r"^executive\s+summary(?:\s*[:\.])?$", s, re.I))


def _match_standalone_heading(line: str) -> str | None:
    s = _normalize_heading_line(line)
    if len(s) > 200:
        return None
    # Executive Summary 与 Abstract 等价（输出仍用 Abstract 节样式时可统一为 abstract）
    if _match_executive_summary_heading(line):
        return "abstract"
    for key, pat in _HEADING_PATTERNS:
        if pat.match(s):
            return key
    if re.match(r"^abstract(?:\s*[:\.])?$", s, re.I):
        return "abstract"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?introduction(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "introduction"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?(?:related\s+work|literature\s+review)(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "related_work"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?(?:methodology|method)(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "method"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?discussion(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "discussion"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?conclusions?(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "conclusion"
    if re.match(
        r"^(?:[IVX]{1,8}\.\s*|\d+\.\s*|\d+\s+)?(?:references?|bibliography|works\s+cited)(?:\s*[:\.])?$",
        s,
        re.I,
    ):
        return "references"
    return None


def _is_front_matter_break_line(line: str) -> bool:
    """标题与摘要之间的分界：出现摘要/关键词等节时停止收集作者行。"""
    if _match_standalone_heading(line) is not None:
        return True
    if _match_executive_summary_heading(line):
        return True
    if _match_keywords_line(line):
        return True
    return False


def _match_keywords_line(line: str) -> bool:
    """Keywords 行：以 Keywords 开头，后接 em dash、连字符或冒号等。"""
    s = _normalize_heading_line(line)
    if len(s) > 500:
        return False
    return bool(re.match(r"^keywords\s*([—\-:]|$)", s, re.I))


def _looks_like_meta_line(line: str) -> bool:
    """作者/单位/日期：用单词边界，避免 'independent' 误匹配 'institute'。"""
    low = line.lower().strip()
    if "@" in line:
        return True
    if re.search(
        r"\b(university|institute|college|laboratory|department|researcher|research|independent|singapore|china|usa|uk)\b",
        low,
    ):
        return True
    if re.search(r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\b", low):
        return True
    if re.search(r"\d{4}", line) and len(line) < 120:
        return True
    if "," in line and len(line) < 200 and not line.endswith("."):
        return True
    return False


def _looks_like_person_name_line(line: str) -> bool:
    """简短姓名行：如 Mingwei Zhang（2–5 个词，无句号）。"""
    s = line.strip()
    if not s or len(s) > 120 or s.endswith("."):
        return False
    parts = s.split()
    if len(parts) < 2 or len(parts) > 6:
        return False
    return bool(re.match(r"^[\w\s\-\.]+$", s, re.UNICODE))


def _collect_title_and_meta(text_lines: list[str]) -> tuple[str | None, list[str]]:
    """
    首段为标题；其后连续若干行为作者区（姓名、单位、日期），参考：
    Mingwei Zhang / Independent Researcher, Singapore / March 2026
    遇到 Abstract、Executive Summary、Keywords 等节起始即停止，不再把节名误当作者行。
    """
    if not text_lines:
        return None, []
    title_text = text_lines[0]
    meta_lines: list[str] = []
    max_meta = 5
    i = 1
    while i < len(text_lines) and len(meta_lines) < max_meta:
        ln = text_lines[i]
        if _is_front_matter_break_line(ln):
            break
        # 作者行：元信息启发式，或像姓名的短行（避免把长摘要句收进来）
        if _looks_like_meta_line(ln) or _looks_like_person_name_line(ln):
            meta_lines.append(ln)
            i += 1
            continue
        # 极短行可能是副标题，但若不像作者则停止
        if len(ln) < 60 and not ln.endswith(".") and len(ln.split()) <= 8:
            meta_lines.append(ln)
            i += 1
            continue
        break
    return title_text, meta_lines


def _add_title(doc: Document, text: str, profile: StyleProfile) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    make_run_font(run, profile.body_font_name, profile.title_size_pt)
    run.bold = True


def _add_meta_line(doc: Document, text: str, profile: StyleProfile) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    make_run_font(run, profile.body_font_name, max(profile.body_font_size_pt - 0.5, 10.0))


def _add_keywords_paragraph(doc: Document, text: str, profile: StyleProfile) -> None:
    """Keywords 行：Keywords 用斜体，其余正文；整段单栏、左对齐。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = profile.line_spacing
    pf.space_after = Pt(8)
    s = text.strip()
    m = re.match(r"^(keywords)(\s*[—\-:]\s*)(.*)$", s, re.I | re.DOTALL)
    if m:
        r1 = p.add_run(m.group(1))
        r1.italic = True
        make_run_font(r1, profile.body_font_name, profile.body_font_size_pt)
        r2 = p.add_run(m.group(2) + m.group(3))
        make_run_font(r2, profile.body_font_name, profile.body_font_size_pt)
    else:
        run = p.add_run(text)
        make_run_font(run, profile.body_font_name, profile.body_font_size_pt)


def _add_section_heading(doc: Document, section_key: str, profile: StyleProfile) -> None:
    """章节标题：加粗 + 略大字号（与正文区分）。"""
    label = _SECTION_DISPLAY.get(section_key, section_key.replace("_", " ").title())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    make_run_font(run, profile.heading_font_name, profile.h1_size_pt)
    run.bold = True


def _add_body_paragraph(doc: Document, text: str, profile: StyleProfile) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = profile.line_spacing
    pf.space_after = Pt(6)
    run = p.add_run(text)
    make_run_font(run, profile.body_font_name, profile.body_font_size_pt)


def _add_reference_line(doc: Document, text: str, profile: StyleProfile) -> None:
    """参考文献区：悬挂缩进，略小字号。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = profile.line_spacing
    pf.left_indent = Pt(18)
    pf.first_line_indent = Pt(-18)
    run = p.add_run(text)
    make_run_font(run, profile.body_font_name, profile.body_font_size_pt - 0.5)


def _has_body_after_title(blocks: list[tuple[str, object]], skip_paragraphs: int) -> bool:
    non_empty_count = 0
    for kind, obj in blocks:
        if kind == "t":
            if non_empty_count >= skip_paragraphs:
                return True
            continue
        t = paragraph_text(obj)  # type: ignore[arg-type]
        if not t:
            continue
        non_empty_count += 1
        if non_empty_count > skip_paragraphs:
            return True
    return False


def _author_lines_from_env() -> list[str]:
    """当源文件未解析出作者行时，用 .env 中的 AUTHOR_LINE_* 注入（个人固定署名）。"""
    out: list[str] = []
    for x in (AUTHOR_LINE_1, AUTHOR_LINE_2, AUTHOR_LINE_3):
        if x and str(x).strip():
            out.append(str(x).strip())
    return out


def _abstract_heading_exists_after_skip(blocks: list[tuple[str, object]], skip_paragraphs: int) -> bool:
    non_empty_count = 0
    for kind, obj in blocks:
        if kind != "p":
            continue
        t = paragraph_text(obj)  # type: ignore[arg-type]
        if not t:
            continue
        non_empty_count += 1
        if non_empty_count <= skip_paragraphs:
            continue
        h = _match_standalone_heading(t)
        if h == "abstract":
            return True
    return False


def format_docx_to_path(
    source_path: Path,
    out_path: Path,
    profile: StyleProfile,
) -> None:
    src = Document(str(source_path))
    blocks = iter_body_blocks(src)
    text_lines = blocks_to_plain_lines(blocks)
    title_text, source_meta_lines = _collect_title_and_meta(text_lines)
    # 仅按源稿段落计数跳过；注入的作者行不在源稿中，不得加大 skip
    skip_paragraphs = (1 if title_text else 0) + len(source_meta_lines)
    display_meta_lines = source_meta_lines if source_meta_lines else _author_lines_from_env()

    out = Document()
    apply_section_margins(out, profile)

    has_title_block = bool(title_text) or bool(display_meta_lines)
    body_cols = max(1, profile.body_columns)
    space_twips = profile.column_spacing_twips if body_cols >= 2 else None

    has_abstract_in_doc = _abstract_heading_exists_after_skip(blocks, skip_paragraphs)

    use_split_after_abstract = (
        body_cols >= 2
        and has_title_block
        and has_abstract_in_doc
    )

    use_split_after_title_meta = (
        body_cols >= 2
        and has_title_block
        and (not has_abstract_in_doc)
        and _has_body_after_title(blocks, skip_paragraphs)
    )

    if use_split_after_abstract or use_split_after_title_meta:
        set_section_column_count(out.sections[0], 1, None)
    elif body_cols >= 2:
        set_section_column_count(out.sections[0], body_cols, space_twips)

    if title_text:
        _add_title(out, title_text, profile)
    for m in display_meta_lines:
        _add_meta_line(out, m, profile)

    if use_split_after_title_meta:
        body_section = out.add_section(WD_SECTION.CONTINUOUS)
        set_section_column_count(body_section, body_cols, space_twips)

    current_section = "body"
    non_empty_para_index = 0
    seen_abstract_heading = False
    abstract_columns_done = use_split_after_title_meta
    post_conclusion_single_done = False

    for kind, obj in blocks:
        if kind == "t":
            append_table_copy(out, obj)  # type: ignore[arg-type]
            continue

        text = paragraph_text(obj)  # type: ignore[arg-type]
        if not text:
            continue

        non_empty_para_index += 1
        if non_empty_para_index <= skip_paragraphs:
            continue

        # Keywords 段：在摘要单栏区内，输出后接双栏（参考排版：Keywords 后再分栏）
        if (
            use_split_after_abstract
            and (not abstract_columns_done)
            and seen_abstract_heading
            and _match_keywords_line(text)
        ):
            _add_keywords_paragraph(out, text, profile)
            body_section = out.add_section(WD_SECTION.CONTINUOUS)
            set_section_column_count(body_section, body_cols, space_twips)
            abstract_columns_done = True
            continue

        head = _match_standalone_heading(text)
        if head:
            # 摘要后双栏：若无 Keywords 段，则在摘要后第一个非 abstract 标题前分栏
            if (
                use_split_after_abstract
                and (not abstract_columns_done)
                and seen_abstract_heading
                and head != "abstract"
            ):
                body_section = out.add_section(WD_SECTION.CONTINUOUS)
                set_section_column_count(body_section, body_cols, space_twips)
                abstract_columns_done = True

            if body_cols >= 2 and (not post_conclusion_single_done):
                split_to_single = False
                if current_section == "conclusion" and head != "conclusion":
                    split_to_single = True
                elif head == "references":
                    split_to_single = True
                if split_to_single:
                    out.add_section(WD_SECTION.CONTINUOUS)
                    set_section_column_count(out.sections[-1], 1, None)
                    post_conclusion_single_done = True

            if head == "abstract":
                seen_abstract_heading = True

            current_section = head
            _add_section_heading(out, head, profile)
            continue

        if current_section == "references":
            _add_reference_line(out, text, profile)
        else:
            _add_body_paragraph(out, text, profile)

    out.save(str(out_path))
    logger.info(
        "Formatted DOCX written to %s (split_abstract=%s split_title_meta=%s has_abstract=%s cols=%s)",
        out_path,
        use_split_after_abstract,
        use_split_after_title_meta,
        has_abstract_in_doc,
        body_cols,
    )
